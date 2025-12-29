"""
Export OpenAPI specification to PDF and Word document formats.

Supports:
- PDF export with formatted documentation
- Word (.docx) export with formatted documentation
"""

import os
import json
import yaml
from typing import Dict, Any, Optional, Tuple
from pathlib import Path


class PDFWordExporter:
    """Export OpenAPI specification to PDF and Word formats."""
    
    def __init__(self, openapi_spec: Dict[str, Any]):
        """Initialize exporter with OpenAPI specification.
        
        Args:
            openapi_spec: OpenAPI specification dictionary
        """
        if not isinstance(openapi_spec, dict):
            raise ValueError(f"openapi_spec must be a dictionary, got {type(openapi_spec)}")
        
        self.openapi_spec = openapi_spec
        self.info = openapi_spec.get('info', {})
        self.servers = openapi_spec.get('servers', [])
        self.paths = openapi_spec.get('paths', {})
        self.components = openapi_spec.get('components', {})
    
    def export_to_pdf(self, output_file: str) -> Tuple[bool, Optional[str]]:
        """Export OpenAPI spec to PDF format.
        
        Args:
            output_file: Path to output PDF file
            
        Returns:
            Tuple of (success: bool, fallback_file: Optional[str])
            If PDF libraries aren't available, returns (False, html_file_path)
        """
        try:
            # On Windows, prefer reportlab (no external DLL dependencies)
            # Try reportlab first (works better on Windows)
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib import colors
                self._export_to_pdf_reportlab(output_file)
                return (True, None)
            except ImportError:
                pass
            except Exception as e:
                # ReportLab import succeeded but execution failed, try weasyprint
                import sys
                if sys.platform == 'win32':
                    # On Windows, if reportlab fails, don't try weasyprint (it has DLL issues)
                    raise Exception(f"ReportLab failed: {str(e)}. On Windows, WeasyPrint requires GTK+ libraries which are difficult to install.")
            
            # Try using weasyprint (better HTML to PDF conversion, but has DLL issues on Windows)
            try:
                from weasyprint import HTML, CSS
                self._export_to_pdf_weasyprint(output_file)
                return (True, None)
            except ImportError:
                pass
            except Exception as e:
                # WeasyPrint import succeeded but execution failed (likely DLL issue on Windows)
                import sys
                if sys.platform == 'win32':
                    # On Windows, weasyprint often fails due to missing GTK+ DLLs
                    # Fall back to HTML
                    html_file = self._export_to_pdf_html_fallback(output_file)
                    return (False, html_file)
                else:
                    raise
            
            # If neither is available, create HTML and return it as fallback
            html_file = self._export_to_pdf_html_fallback(output_file)
            return (False, html_file)
        except Exception as e:
            # If it's a known DLL/library loading error, provide HTML fallback
            error_str = str(e).lower()
            if 'dll' in error_str or 'library' in error_str or 'cannot load' in error_str:
                html_file = self._export_to_pdf_html_fallback(output_file)
                return (False, html_file)
            raise Exception(f"Failed to export to PDF: {str(e)}")
    
    def _export_to_pdf_weasyprint(self, output_file: str) -> bool:
        """Export using WeasyPrint."""
        from weasyprint import HTML, CSS
        
        html_content = self._generate_html_content()
        
        HTML(string=html_content).write_pdf(output_file)
        return True
    
    def _export_to_pdf_reportlab(self, output_file: str) -> bool:
        """Export using ReportLab."""
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate(output_file, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#007acc'),
            spaceAfter=30,
        )
        title = self.info.get('title', 'API Documentation')
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Description
        if self.info.get('description'):
            desc_style = ParagraphStyle(
                'Description',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
            )
            story.append(Paragraph(self.info.get('description', ''), desc_style))
            story.append(Spacer(1, 0.2*inch))
        
        # API Information
        info_data = [
            ['Version', self.info.get('version', 'N/A')],
            ['OpenAPI Version', self.openapi_spec.get('openapi', 'N/A')],
        ]
        if self.info.get('contact', {}).get('email'):
            info_data.append(['Contact', self.info['contact']['email']])
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Servers
        if self.servers:
            story.append(Paragraph('<b>Base URLs</b>', styles['Heading2']))
            for server in self.servers:
                server_url = server.get('url', '')
                server_desc = server.get('description', '')
                story.append(Paragraph(f"• {server_url}", styles['Normal']))
                if server_desc:
                    story.append(Paragraph(f"  <i>{server_desc}</i>", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Paths
        story.append(PageBreak())
        story.append(Paragraph('<b>API Endpoints</b>', styles['Heading1']))
        story.append(Spacer(1, 0.2*inch))
        
        for path, path_item in self.paths.items():
            story.append(Paragraph(f"<b>{path}</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            for method, operation in path_item.items():
                if method.lower() not in ['get', 'post', 'put', 'patch', 'delete', 'head', 'options']:
                    continue
                
                method_upper = method.upper()
                operation_id = operation.get('operationId', '')
                summary = operation.get('summary', '')
                description = operation.get('description', '')
                tags = operation.get('tags', [])
                
                # Use operation ID only (e.g., "getUserId", "createUser")
                if operation_id:
                    method_text = f"<b>{operation_id}</b>"
                else:
                    # Fallback to just method if no operation ID (no summary)
                    method_text = f"<b>{method_upper}</b>"
                story.append(Paragraph(method_text, styles['Heading3']))
                
                # Tags
                if tags:
                    tags_text = "Tags: " + ", ".join(tags)
                    story.append(Paragraph(f"<i>{tags_text}</i>", styles['Normal']))
                
                # Description
                if description:
                    story.append(Paragraph(description, styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
                
                # Parameters
                parameters = operation.get('parameters', [])
                if parameters:
                    story.append(Paragraph("<b>Parameters:</b>", styles['Normal']))
                    param_data = [['Name', 'In', 'Type', 'Required', 'Description']]
                    for param in parameters:
                        param_name = param.get('name', '')
                        param_in = param.get('in', '')
                        param_schema = param.get('schema', {})
                        param_type = param_schema.get('type', 'string')
                        param_required = 'Yes' if param.get('required', False) else 'No'
                        param_desc = param.get('description', '')
                        param_data.append([param_name, param_in, param_type, param_required, param_desc])
                    
                    param_table = Table(param_data, colWidths=[1.2*inch, 0.8*inch, 0.8*inch, 0.6*inch, 2.6*inch])
                    param_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    story.append(param_table)
                    story.append(Spacer(1, 0.1*inch))
                
                # Request Body
                request_body = operation.get('requestBody')
                if request_body:
                    story.append(Paragraph("<b>Request Body:</b>", styles['Normal']))
                    content = request_body.get('content', {})
                    for content_type, content_spec in content.items():
                        schema = content_spec.get('schema', {})
                        example = content_spec.get('example')
                        
                        story.append(Paragraph(f"Content-Type: <i>{content_type}</i>", styles['Normal']))
                        
                        # Schema information
                        if schema:
                            schema_type = schema.get('type', 'object')
                            schema_ref = schema.get('$ref', '')
                            if schema_ref:
                                # Extract schema name from reference
                                schema_name = schema_ref.split('/')[-1]
                                story.append(Paragraph(f"Schema: <b>{schema_name}</b>", styles['Normal']))
                            else:
                                story.append(Paragraph(f"Type: {schema_type}", styles['Normal']))
                        
                        # Example
                        if example:
                            import json
                            example_str = json.dumps(example, indent=2) if isinstance(example, (dict, list)) else str(example)
                            # Truncate if too long
                            if len(example_str) > 500:
                                example_str = example_str[:500] + "... (truncated)"
                            story.append(Paragraph(f"Example:<br/><font face='Courier' size='8'>{example_str}</font>", styles['Normal']))
                    
                    story.append(Spacer(1, 0.1*inch))
                
                # Responses
                responses = operation.get('responses', {})
                if responses:
                    story.append(Paragraph("<b>Responses:</b>", styles['Normal']))
                    for status_code, response_spec in responses.items():
                        response_desc = response_spec.get('description', '')
                        response_content = response_spec.get('content', {})
                        
                        status_para = f"<b>{status_code}</b>"
                        if response_desc:
                            status_para += f" - {response_desc}"
                        story.append(Paragraph(status_para, styles['Normal']))
                        
                        # Response schema and examples
                        for content_type, content_spec in response_content.items():
                            schema = content_spec.get('schema', {})
                            example = content_spec.get('example')
                            
                            if schema:
                                schema_type = schema.get('type', 'object')
                                schema_ref = schema.get('$ref', '')
                                if schema_ref:
                                    schema_name = schema_ref.split('/')[-1]
                                    story.append(Paragraph(f"  Schema: <b>{schema_name}</b> ({content_type})", styles['Normal']))
                                else:
                                    story.append(Paragraph(f"  Type: {schema_type} ({content_type})", styles['Normal']))
                            
                            if example:
                                import json
                                example_str = json.dumps(example, indent=2) if isinstance(example, (dict, list)) else str(example)
                                if len(example_str) > 300:
                                    example_str = example_str[:300] + "... (truncated)"
                                story.append(Paragraph(f"  Example:<br/><font face='Courier' size='7'>{example_str}</font>", styles['Normal']))
                    
                    story.append(Spacer(1, 0.1*inch))
                
                # Security
                security = operation.get('security', [])
                if security:
                    story.append(Paragraph("<b>Security:</b>", styles['Normal']))
                    for sec_req in security:
                        for scheme_name, scopes in sec_req.items():
                            scopes_str = ", ".join(scopes) if scopes else "No scopes"
                            story.append(Paragraph(f"  <b>{scheme_name}</b>: {scopes_str}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                
                story.append(Spacer(1, 0.15*inch))
        
        # Components Schemas
        schemas = self.components.get('schemas', {}).copy()
        
        # Extract inline nested objects from all schemas and add them to schemas dict
        extracted_schemas = {}
        # Map to track which extracted schema corresponds to which inline object path
        extracted_schema_map = {}  # Maps (parent_name, prop_name) -> schema_name
        
        for schema_name, schema in schemas.items():
            self._extract_nested_schemas(schema, extracted_schemas, extracted_schema_map, parent_name=schema_name)
        
        # Merge extracted schemas into main schemas dict
        schemas.update(extracted_schemas)
        
        if schemas:
            story.append(PageBreak())
            story.append(Paragraph('<b>Data Models</b>', styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
            
            # Sort schemas: root schemas first, then nested schemas
            root_schemas = {k: v for k, v in schemas.items() if k not in extracted_schemas}
            for schema_name, schema in root_schemas.items():
                story.append(Paragraph(f"<b>{schema_name}</b>", styles['Heading2']))
                
                schema_type = schema.get('type', 'object')
                schema_desc = schema.get('description', '')
                
                if schema_desc:
                    story.append(Paragraph(schema_desc, styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
                
                story.append(Paragraph(f"Type: <b>{schema_type}</b>", styles['Normal']))
                
                # Properties
                properties = schema.get('properties', {})
                required_fields = schema.get('required', [])
                
                if properties:
                    story.append(Paragraph("<b>Properties:</b>", styles['Normal']))
                    prop_data = [['Property', 'Type', 'Required', 'Description']]
                    for prop_name, prop_schema in properties.items():
                        prop_ref = prop_schema.get('$ref', '')
                        prop_type = prop_schema.get('type', '')
                        
                        # Check for $ref first (schema reference)
                        if prop_ref:
                            # Reference to another schema - show schema name
                            prop_type = prop_ref.split('/')[-1]
                        elif prop_schema.get('items'):
                            # Array type
                            items = prop_schema.get('items', {})
                            items_type = items.get('type', 'object')
                            items_ref = items.get('$ref', '')
                            if items_ref:
                                prop_type = f"array[{items_ref.split('/')[-1]}]"
                            elif items_type == 'object' and items.get('properties'):
                                # Inline nested object in array - check if it was extracted
                                extracted_name = extracted_schema_map.get((schema_name, f"{prop_name}Item"))
                                if extracted_name:
                                    prop_type = f"array[{extracted_name}]"
                                else:
                                    nested_props_count = len(items.get('properties', {}))
                                    prop_type = f"array[object ({nested_props_count} properties)]"
                            else:
                                prop_type = f"array[{items_type}]"
                        elif prop_type == 'object' or (not prop_type and prop_schema.get('properties')):
                            # Inline nested object - check if it was extracted as a separate schema
                            nested_props = prop_schema.get('properties', {})
                            if nested_props:
                                # Check if this inline object was extracted as a separate schema
                                extracted_name = extracted_schema_map.get((schema_name, prop_name))
                                if extracted_name:
                                    prop_type = extracted_name
                                else:
                                    # Fallback: show object with properties count
                                    nested_props_count = len(nested_props)
                                    prop_type = f"object ({nested_props_count} properties)"
                            else:
                                # No properties, just show as object
                                prop_type = 'object'
                        elif not prop_type:
                            # Fallback if no type specified
                            prop_type = 'string'
                        
                        prop_required = 'Yes' if prop_name in required_fields else 'No'
                        prop_desc = prop_schema.get('description', '')
                        prop_data.append([prop_name, prop_type, prop_required, prop_desc])
                    
                    prop_table = Table(prop_data, colWidths=[1.5*inch, 1.2*inch, 0.8*inch, 2.5*inch])
                    prop_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    story.append(prop_table)
                    story.append(Spacer(1, 0.1*inch))
                
                # Example
                example = schema.get('example')
                if example:
                    import json
                    example_str = json.dumps(example, indent=2) if isinstance(example, (dict, list)) else str(example)
                    if len(example_str) > 400:
                        example_str = example_str[:400] + "... (truncated)"
                    story.append(Paragraph(f"Example:<br/><font face='Courier' size='8'>{example_str}</font>", styles['Normal']))
                
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        return True
    
    def _export_to_pdf_html_fallback(self, output_file: str) -> str:
        """Fallback: Generate HTML that can be converted to PDF.
        
        Returns:
            Path to the generated HTML file
        """
        html_content = self._generate_html_content()
        html_file = output_file.replace('.pdf', '.html')
        
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return html_file
    
    def export_to_word(self, output_file: str) -> bool:
        """Export OpenAPI spec to Word (.docx) format.
        
        Args:
            output_file: Path to output Word file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise Exception(
                "python-docx library not installed. "
                "Install it with: pip install python-docx"
            )
        
        doc = Document()
        
        # Title
        title = doc.add_heading(self.info.get('title', 'API Documentation'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Description
        if self.info.get('description'):
            doc.add_paragraph(self.info.get('description', ''))
            doc.add_paragraph()
        
        # API Information
        info_heading = doc.add_heading('API Information', level=1)
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Version', self.info.get('version', 'N/A')),
            ('OpenAPI Version', self.openapi_spec.get('openapi', 'N/A')),
        ]
        if self.info.get('contact', {}).get('email'):
            info_data.append(('Contact', self.info['contact']['email']))
        
        for label, value in info_data:
            row = info_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Servers
        if self.servers:
            doc.add_heading('Base URLs', level=1)
            for server in self.servers:
                server_url = server.get('url', '')
                server_desc = server.get('description', '')
                p = doc.add_paragraph(server_url, style='List Bullet')
                if server_desc:
                    doc.add_paragraph(server_desc, style='Intense Quote')
            doc.add_paragraph()
        
        # Paths
        doc.add_heading('API Endpoints', level=1)
        
        for path, path_item in self.paths.items():
            doc.add_heading(path, level=2)
            
            for method, operation in path_item.items():
                if method.lower() not in ['get', 'post', 'put', 'patch', 'delete', 'head', 'options']:
                    continue
                
                method_upper = method.upper()
                operation_id = operation.get('operationId', '')
                summary = operation.get('summary', '')
                description = operation.get('description', '')
                tags = operation.get('tags', [])
                
                # Use operation ID only (e.g., "getUserId", "createUser")
                method_para = doc.add_paragraph()
                if operation_id:
                    method_text = operation_id
                else:
                    # Fallback to just method if no operation ID (no summary)
                    method_text = f"{method_upper}"
                method_run = method_para.add_run(method_text)
                method_run.bold = True
                method_run.font.size = Pt(12)
                
                # Tags
                if tags:
                    tags_para = doc.add_paragraph()
                    tags_run = tags_para.add_run(f"Tags: {', '.join(tags)}")
                    tags_run.italic = True
                    tags_run.font.size = Pt(10)
                
                # Description
                if description:
                    doc.add_paragraph(description)
                    doc.add_paragraph()
                
                # Parameters
                parameters = operation.get('parameters', [])
                if parameters:
                    doc.add_heading('Parameters', level=3)
                    param_table = doc.add_table(rows=1, cols=5)
                    param_table.style = 'Light Grid Accent 1'
                    param_table.rows[0].cells[0].text = 'Name'
                    param_table.rows[0].cells[1].text = 'In'
                    param_table.rows[0].cells[2].text = 'Type'
                    param_table.rows[0].cells[3].text = 'Required'
                    param_table.rows[0].cells[4].text = 'Description'
                    
                    for param in parameters:
                        row = param_table.add_row()
                        row.cells[0].text = param.get('name', '')
                        row.cells[1].text = param.get('in', '')
                        schema = param.get('schema', {})
                        param_type = schema.get('type', 'string')
                        param_ref = schema.get('$ref', '')
                        if param_ref:
                            param_type = param_ref.split('/')[-1]
                        row.cells[2].text = param_type
                        row.cells[3].text = 'Yes' if param.get('required', False) else 'No'
                        row.cells[4].text = param.get('description', '')
                    doc.add_paragraph()
                
                # Request Body
                request_body = operation.get('requestBody')
                if request_body:
                    doc.add_heading('Request Body', level=3)
                    content = request_body.get('content', {})
                    for content_type, content_spec in content.items():
                        schema = content_spec.get('schema', {})
                        example = content_spec.get('example')
                        
                        type_para = doc.add_paragraph()
                        type_run = type_para.add_run(f"Content-Type: {content_type}")
                        type_run.italic = True
                        
                        if schema:
                            schema_type = schema.get('type', 'object')
                            schema_ref = schema.get('$ref', '')
                            if schema_ref:
                                schema_name = schema_ref.split('/')[-1]
                                doc.add_paragraph(f"Schema: {schema_name}")
                            else:
                                doc.add_paragraph(f"Type: {schema_type}")
                        
                        if example:
                            import json
                            example_str = json.dumps(example, indent=2) if isinstance(example, (dict, list)) else str(example)
                            if len(example_str) > 500:
                                example_str = example_str[:500] + "... (truncated)"
                            example_para = doc.add_paragraph('Example:')
                            example_code = doc.add_paragraph(example_str)
                            example_code.style = 'No Spacing'
                            for run in example_code.runs:
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                    doc.add_paragraph()
                
                # Responses
                responses = operation.get('responses', {})
                if responses:
                    doc.add_heading('Responses', level=3)
                    for status_code, response_spec in responses.items():
                        response_desc = response_spec.get('description', '')
                        response_content = response_spec.get('content', {})
                        
                        status_para = doc.add_paragraph()
                        status_run = status_para.add_run(f"{status_code}")
                        status_run.bold = True
                        if response_desc:
                            status_para.add_run(f" - {response_desc}")
                        
                        for content_type, content_spec in response_content.items():
                            schema = content_spec.get('schema', {})
                            example = content_spec.get('example')
                            
                            if schema:
                                schema_type = schema.get('type', 'object')
                                schema_ref = schema.get('$ref', '')
                                if schema_ref:
                                    schema_name = schema_ref.split('/')[-1]
                                    doc.add_paragraph(f"  Schema: {schema_name} ({content_type})", style='List Bullet 2')
                                else:
                                    doc.add_paragraph(f"  Type: {schema_type} ({content_type})", style='List Bullet 2')
                            
                            if example:
                                import json
                                example_str = json.dumps(example, indent=2) if isinstance(example, (dict, list)) else str(example)
                                if len(example_str) > 300:
                                    example_str = example_str[:300] + "... (truncated)"
                                example_para = doc.add_paragraph('  Example:', style='List Bullet 2')
                                example_code = doc.add_paragraph(example_str)
                                example_code.style = 'No Spacing'
                                for run in example_code.runs:
                                    run.font.name = 'Courier New'
                                    run.font.size = Pt(8)
                    doc.add_paragraph()
                
                # Security
                security = operation.get('security', [])
                if security:
                    doc.add_heading('Security', level=3)
                    for sec_req in security:
                        for scheme_name, scopes in sec_req.items():
                            scopes_str = ", ".join(scopes) if scopes else "No scopes"
                            doc.add_paragraph(f"{scheme_name}: {scopes_str}", style='List Bullet')
                    doc.add_paragraph()
                
                doc.add_paragraph()
        
        # Components Schemas
        schemas = self.components.get('schemas', {}).copy()
        
        # Extract inline nested objects from all schemas and add them to schemas dict
        extracted_schemas = {}
        extracted_schema_map = {}  # Maps (parent_name, prop_name) -> schema_name
        
        for schema_name, schema in schemas.items():
            self._extract_nested_schemas(schema, extracted_schemas, extracted_schema_map, parent_name=schema_name)
        
        # Merge extracted schemas into main schemas dict
        schemas.update(extracted_schemas)
        
        if schemas:
            doc.add_page_break()
            doc.add_heading('Data Models', level=1)
            
            # Sort schemas: root schemas first, then nested schemas
            root_schemas = {k: v for k, v in schemas.items() if k not in extracted_schemas}
            for schema_name, schema in root_schemas.items():
                doc.add_heading(schema_name, level=2)
                
                schema_type = schema.get('type', 'object')
                schema_desc = schema.get('description', '')
                
                if schema_desc:
                    doc.add_paragraph(schema_desc)
                
                doc.add_paragraph(f"Type: {schema_type}")
                
                # Properties
                properties = schema.get('properties', {})
                required_fields = schema.get('required', [])
                
                if properties:
                    doc.add_heading('Properties', level=3)
                    prop_table = doc.add_table(rows=1, cols=4)
                    prop_table.style = 'Light Grid Accent 1'
                    prop_table.rows[0].cells[0].text = 'Property'
                    prop_table.rows[0].cells[1].text = 'Type'
                    prop_table.rows[0].cells[2].text = 'Required'
                    prop_table.rows[0].cells[3].text = 'Description'
                    
                    for prop_name, prop_schema in properties.items():
                        row = prop_table.add_row()
                        row.cells[0].text = prop_name
                        
                        prop_ref = prop_schema.get('$ref', '')
                        prop_type = prop_schema.get('type', '')
                        
                        # Check for $ref first (schema reference)
                        if prop_ref:
                            # Reference to another schema - show schema name
                            prop_type = prop_ref.split('/')[-1]
                        elif prop_schema.get('items'):
                            # Array type
                            items = prop_schema.get('items', {})
                            items_type = items.get('type', 'object')
                            items_ref = items.get('$ref', '')
                            if items_ref:
                                prop_type = f"array[{items_ref.split('/')[-1]}]"
                            elif items_type == 'object' and items.get('properties'):
                                # Inline nested object in array - check if it was extracted
                                extracted_name = extracted_schema_map.get((schema_name, f"{prop_name}Item"))
                                if extracted_name:
                                    prop_type = f"array[{extracted_name}]"
                                else:
                                    nested_props_count = len(items.get('properties', {}))
                                    prop_type = f"array[object ({nested_props_count} properties)]"
                            else:
                                prop_type = f"array[{items_type}]"
                        elif prop_type == 'object' or (not prop_type and prop_schema.get('properties')):
                            # Inline nested object - check if it was extracted as a separate schema
                            nested_props = prop_schema.get('properties', {})
                            if nested_props:
                                # Check if this inline object was extracted as a separate schema
                                extracted_name = extracted_schema_map.get((schema_name, prop_name))
                                if extracted_name:
                                    prop_type = extracted_name
                                else:
                                    # Fallback: show object with properties count
                                    nested_props_count = len(nested_props)
                                    prop_type = f"object ({nested_props_count} properties)"
                            else:
                                # No properties, just show as object
                                prop_type = 'object'
                        elif not prop_type:
                            # Fallback if no type specified
                            prop_type = 'string'
                        
                        row.cells[1].text = prop_type
                        row.cells[2].text = 'Yes' if prop_name in required_fields else 'No'
                        row.cells[3].text = prop_schema.get('description', '')
                    doc.add_paragraph()
                
                # Example
                example = schema.get('example')
                if example:
                    import json
                    example_str = json.dumps(example, indent=2) if isinstance(example, (dict, list)) else str(example)
                    if len(example_str) > 400:
                        example_str = example_str[:400] + "... (truncated)"
                    doc.add_paragraph('Example:')
                    example_code = doc.add_paragraph(example_str)
                    example_code.style = 'No Spacing'
                    for run in example_code.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                
                doc.add_paragraph()
            
            # Now display extracted nested schemas
            if extracted_schemas:
                doc.add_paragraph()
                doc.add_heading('Nested Object Schemas', level=2)
                doc.add_paragraph()
                
                for schema_name, schema in sorted(extracted_schemas.items()):
                    doc.add_heading(schema_name, level=3)
                    
                    schema_type = schema.get('type', 'object')
                    schema_desc = schema.get('description', '')
                    
                    if schema_desc:
                        doc.add_paragraph(schema_desc)
                    
                    doc.add_paragraph(f"Type: {schema_type}")
                    
                    # Properties
                    properties = schema.get('properties', {})
                    required_fields = schema.get('required', [])
                    
                    if properties:
                        doc.add_heading('Properties', level=4)
                        prop_table = doc.add_table(rows=1, cols=4)
                        prop_table.style = 'Light Grid Accent 1'
                        prop_table.rows[0].cells[0].text = 'Property'
                        prop_table.rows[0].cells[1].text = 'Type'
                        prop_table.rows[0].cells[2].text = 'Required'
                        prop_table.rows[0].cells[3].text = 'Description'
                        
                        for prop_name, prop_schema in properties.items():
                            row = prop_table.add_row()
                            row.cells[0].text = prop_name
                            
                            prop_ref = prop_schema.get('$ref', '')
                            prop_type = prop_schema.get('type', '')
                            
                            if prop_ref:
                                prop_type = prop_ref.split('/')[-1]
                            elif prop_schema.get('items'):
                                items = prop_schema.get('items', {})
                                items_type = items.get('type', 'object')
                                items_ref = items.get('$ref', '')
                                if items_ref:
                                    prop_type = f"array[{items_ref.split('/')[-1]}]"
                                elif items_type == 'object' and items.get('properties'):
                                    # Inline nested object in array - check if it was extracted
                                    extracted_name = extracted_schema_map.get((schema_name, f"{prop_name}Item"))
                                    if extracted_name:
                                        prop_type = f"array[{extracted_name}]"
                                    else:
                                        nested_props_count = len(items.get('properties', {}))
                                        prop_type = f"array[object ({nested_props_count} properties)]"
                                else:
                                    prop_type = f"array[{items_type}]"
                            elif prop_type == 'object' or (not prop_type and prop_schema.get('properties')):
                                # Inline nested object - check if it was extracted as a separate schema
                                nested_props = prop_schema.get('properties', {})
                                if nested_props:
                                    # Check if this inline object was extracted as a separate schema
                                    extracted_name = extracted_schema_map.get((schema_name, prop_name))
                                    if extracted_name:
                                        prop_type = extracted_name
                                    else:
                                        nested_props_count = len(nested_props)
                                        prop_type = f"object ({nested_props_count} properties)"
                                else:
                                    prop_type = 'object'
                            elif not prop_type:
                                # Fallback if no type specified
                                prop_type = 'string'
                            
                            row.cells[1].text = prop_type
                            row.cells[2].text = 'Yes' if prop_name in required_fields else 'No'
                            row.cells[3].text = prop_schema.get('description', '')
                        doc.add_paragraph()
                    
                    doc.add_paragraph()
        
        doc.save(output_file)
        return True
    
    def _extract_nested_schemas(self, schema: Dict[str, Any], extracted_schemas: Dict[str, Any], extracted_schema_map: Dict[tuple, str], parent_name: str = "", counter: int = 1):
        """Recursively extract inline nested objects from a schema.
        
        Args:
            schema: Schema dictionary to extract from
            extracted_schemas: Dictionary to store extracted schemas
            extracted_schema_map: Dictionary mapping (parent_name, prop_name) -> schema_name
            parent_name: Name of parent schema (for naming nested schemas)
            counter: Counter for unique naming
        """
        if not isinstance(schema, dict):
            return
        
        properties = schema.get('properties', {})
        for prop_name, prop_schema in properties.items():
            prop_ref = prop_schema.get('$ref', '')
            prop_type = prop_schema.get('type', '')
            
            # If it's an inline nested object (not a reference)
            if not prop_ref and (prop_type == 'object' or (not prop_type and prop_schema.get('properties'))):
                nested_props = prop_schema.get('properties', {})
                if nested_props:
                    # Generate a schema name
                    schema_name = self._generate_nested_schema_name(prop_name, parent_name, counter)
                    counter += 1
                    
                    # Store the nested schema
                    extracted_schemas[schema_name] = prop_schema.copy()
                    
                    # Map this inline object to its extracted schema name
                    extracted_schema_map[(parent_name, prop_name)] = schema_name
                    
                    # Recursively extract nested objects from this nested schema
                    self._extract_nested_schemas(prop_schema, extracted_schemas, extracted_schema_map, schema_name, counter)
            
            # Check arrays with nested objects
            elif prop_schema.get('items'):
                items = prop_schema.get('items', {})
                items_ref = items.get('$ref', '')
                items_type = items.get('type', 'object')
                
                if not items_ref and items_type == 'object' and items.get('properties'):
                    # Generate a schema name for array item
                    schema_name = self._generate_nested_schema_name(prop_name, parent_name, counter, is_array=True)
                    counter += 1
                    
                    # Store the nested schema
                    extracted_schemas[schema_name] = items.copy()
                    
                    # Map this array item to its extracted schema name
                    extracted_schema_map[(parent_name, f"{prop_name}Item")] = schema_name
                    
                    # Recursively extract nested objects
                    self._extract_nested_schemas(items, extracted_schemas, extracted_schema_map, schema_name, counter)
    
    def _generate_nested_schema_name(self, prop_name: str, parent_name: str, counter: int, is_array: bool = False) -> str:
        """Generate a name for a nested schema.
        
        Args:
            prop_name: Property name (e.g., 'payment_method')
            parent_name: Parent schema name (e.g., 'CreateUserRequest')
            counter: Counter for uniqueness
            is_array: Whether this is for an array item
            
        Returns:
            Schema name (e.g., 'PaymentMethod', 'CreateUserRequestPaymentMethod')
        """
        # Convert property name to PascalCase
        name_parts = prop_name.replace('_', ' ').title().replace(' ', '')
        
        if parent_name:
            # Try to create a meaningful name
            base_name = parent_name.replace('Request', '').replace('Response', '')
            schema_name = f"{base_name}{name_parts}"
        else:
            schema_name = name_parts
        
        if is_array:
            schema_name = f"{schema_name}Item"
        
        # Ensure uniqueness
        existing_schemas = self.components.get('schemas', {})
        if schema_name in existing_schemas:
            schema_name = f"{schema_name}{counter}"
        
        return schema_name
    
    def _generate_html_content(self) -> str:
        """Generate HTML content for PDF export."""
        title = self.info.get('title', 'API Documentation')
        description = self.info.get('description', '')
        version = self.info.get('version', 'N/A')
        openapi_version = self.openapi_spec.get('openapi', 'N/A')
        
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 40px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            color: #007acc;
            border-bottom: 2px solid #007acc;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #005a9e;
            margin-top: 30px;
        }}
        h3 {{
            color: #1a8cd8;
            margin-top: 20px;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #f0f0f0;
            font-weight: bold;
        }}
        .method {{
            display: inline-block;
            padding: 4px 8px;
            border-radius: 4px;
            font-weight: bold;
            margin-right: 10px;
        }}
        .method.get {{ background-color: #61affe; color: white; }}
        .method.post {{ background-color: #49cc90; color: white; }}
        .method.put {{ background-color: #fca130; color: white; }}
        .method.delete {{ background-color: #f93e3e; color: white; }}
        .method.patch {{ background-color: #50e3c2; color: white; }}
        code {{
            background-color: #f5f5f5;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: monospace;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p>{description}</p>
    
    <h2>API Information</h2>
    <table>
        <tr><th>Version</th><td>{version}</td></tr>
        <tr><th>OpenAPI Version</th><td>{openapi_version}</td></tr>
    </table>
"""
        
        # Servers
        if self.servers:
            html += "<h2>Base URLs</h2><ul>"
            for server in self.servers:
                server_url = server.get('url', '')
                server_desc = server.get('description', '')
                html += f"<li><code>{server_url}</code>"
                if server_desc:
                    html += f" - {server_desc}"
                html += "</li>"
            html += "</ul>"
        
        # Paths
        html += "<h2>API Endpoints</h2>"
        for path, path_item in self.paths.items():
            html += f"<h3>{path}</h3>"
            for method, operation in path_item.items():
                if method.lower() not in ['get', 'post', 'put', 'patch', 'delete', 'head', 'options']:
                    continue
                
                method_upper = method.upper()
                operation_id = operation.get('operationId', '')
                summary = operation.get('summary', '')
                description = operation.get('description', '')
                
                # Use operation ID only (e.g., "getUserId", "createUser")
                if operation_id:
                    html += f'<p><span class="method {method.lower()}">{operation_id}</span></p>'
                else:
                    # Fallback to just method if no operation ID
                    html += f'<p><span class="method {method.lower()}">{method_upper}</span></p>'
                if description:
                    html += f"<p>{description}</p>"
        
        html += """
</body>
</html>
"""
        return html

